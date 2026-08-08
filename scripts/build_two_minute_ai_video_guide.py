from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/two-minute-ai-answers-first-10.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1E2630"
MUTED = "667085"
PALE_BLUE = "E8EEF5"
PALE_LIME = "EAF5B5"
PALE_ORANGE = "FBE2D5"
WHITE = "FFFFFF"


VIDEOS = [
    {
        "number": "01",
        "category": "Choosing a tool",
        "question": "Should I use ChatGPT or Claude for my business?",
        "hook": "Do not choose your AI tool from a feature checklist. Choose it from the work you need to finish every week.",
        "answer": "For most small businesses, either tool can handle everyday writing, research, planning, and analysis. The better choice is the one that fits your recurring workflow and that your team will actually use.",
        "points": [
            "Start with one real weekly task: drafting proposals, summarizing documents, planning content, or analyzing customer feedback.",
            "Run the same task in both tools with the same instructions and source material.",
            "Compare the finished output, the amount of correction required, ease of use, integrations, and how comfortable you are with the privacy controls.",
            "Pick one primary tool for 30 days. Add a second only when it solves a specific gap.",
        ],
        "example": "Take a customer email thread and ask both tools to produce a reply, a three-item action list, and the missing questions. The winner is not the most impressive response; it is the one you would confidently use after the fewest edits.",
        "screen": ["Same task", "Same context", "Compare corrections", "Choose for 30 days"],
        "close": "If you send me the task you repeat most often, I will tell you which tool I would test first and why.",
        "caption": "ChatGPT or Claude? Start with the work—not the logo. Run one real business task through both tools using the same prompt and source material. Compare accuracy, editing time, integrations, and whether your team will actually use it. Pick one for 30 days before adding another subscription. What task would you test?",
        "library": "ChatGPT and Claude overlap on many everyday business tasks. The most useful comparison is a side-by-side test using your own recurring work, evaluated by accuracy, editing time, fit, and adoption.",
    },
    {
        "number": "02",
        "category": "Choosing a tool",
        "question": "Which Claude model should I use?",
        "hook": "The best Claude model is usually not the biggest one. It is the least expensive model that completes the task reliably.",
        "answer": "Use Claude's default or balanced model for normal business work. Move to the fastest option for simple, high-volume tasks and the strongest reasoning option for complex decisions, difficult analysis, or work where mistakes are expensive.",
        "points": [
            "Use the balanced/default model for writing, summarizing, planning, and most document work.",
            "Use the fast model for classification, cleanup, short rewrites, and repetitive low-risk work.",
            "Use the strongest reasoning model when the task has ambiguity, many constraints, or needs careful tradeoffs.",
            "Judge the model by the complete workflow: quality, correction time, speed, and usage—not by a benchmark headline.",
        ],
        "example": "A social caption does not need your most powerful model. A decision memo comparing three software contracts probably does. Use the smallest hammer that drives the nail.",
        "screen": ["Fast: simple + repetitive", "Default: everyday work", "Strongest: complex + high-stakes"],
        "close": "Tell me the task and what a mistake would cost you. That is usually enough to choose the right model tier.",
        "caption": "Which Claude model should you use? Fast for simple repetitive work. Default for everyday writing, research, and planning. Strongest reasoning for ambiguity, major tradeoffs, and expensive mistakes. Use the least costly model that produces reliable work after review.",
        "library": "Choose a Claude model based on task complexity and the cost of a mistake. Start with the default model, move down for simple volume work, and move up when deeper reasoning materially improves the result.",
    },
    {
        "number": "03",
        "category": "Choosing a tool",
        "question": "Do I need more than one AI tool?",
        "hook": "Most businesses do not have an AI-tool shortage. They have an unfinished-workflow problem.",
        "answer": "Start with one general AI assistant and make one workflow dependable. Add another tool only when you can name the job it does better, the time it saves, and who will own it.",
        "points": [
            "One primary tool reduces training, duplicated subscriptions, scattered history, and inconsistent processes.",
            "A second tool makes sense when it provides a needed integration, specialized capability, or measurable quality improvement.",
            "Avoid buying a separate AI app for every department before the underlying workflow is clear.",
            "Review the stack quarterly and remove anything nobody uses or nobody owns.",
        ],
        "example": "If Claude handles research, writing, and document work, do not add three writing tools because each has one clever demo. Add a second tool when, for example, it connects directly to a system Claude cannot reach and saves a repeatable hour every week.",
        "screen": ["One primary assistant", "One proven workflow", "Add only for a named gap"],
        "close": "List the AI tools you pay for and the job each one owns. If a tool has no clear job, that is your first clue.",
        "caption": "You probably do not need five AI tools. Start with one general assistant and one dependable workflow. Add another tool only when it fills a named gap, saves measurable time, and has a clear owner. More subscriptions do not automatically create more leverage.",
        "library": "A small business usually benefits more from one well-adopted AI tool than a crowded stack. Add specialized tools only for clear gaps, then review usage and ownership regularly.",
    },
    {
        "number": "04",
        "category": "Choosing a tool",
        "question": "Is a paid AI plan worth it for a small business?",
        "hook": "A paid AI plan is worth it when it replaces more time than it costs—not when the demo looks impressive.",
        "answer": "Run a two-week test around one repeated task. If the paid plan saves useful time, improves quality, or unlocks a capability you will use consistently, keep it. If not, cancel it.",
        "points": [
            "Choose one task that happens at least weekly and measure how long it takes today.",
            "Count setup, review, and correction time—not just generation time.",
            "Track whether the paid features change the result: larger limits, better models, file handling, connectors, or browser work.",
            "Set a simple threshold, such as saving one useful hour per month or preventing one recurring error.",
        ],
        "example": "If a $20-to-$30 plan saves two hours of proposal cleanup each month, the decision is easy. If you only use it to rewrite an occasional sentence, the free tier may be enough.",
        "screen": ["Pick one weekly task", "Measure total time", "Test for 2 weeks", "Keep or cancel"],
        "close": "Do not ask whether AI is worth paying for in general. Ask whether this plan earns its place in one workflow.",
        "caption": "Is a paid AI plan worth it? Test one weekly task for two weeks. Measure the entire workflow: setup, output, corrections, and review. Keep the plan if it saves meaningful time or unlocks a capability you use. Cancel it if the value only appears in the demo.",
        "library": "Evaluate a paid AI plan with a short workflow test, not a feature list. Measure total time saved, output quality, and whether paid-only capabilities are used consistently.",
    },
    {
        "number": "05",
        "category": "Choosing a tool",
        "question": "When should I use Claude Cowork instead of regular Claude?",
        "hook": "Use regular Claude when you need an answer. Use Cowork when you need a bounded piece of work completed across files, tools, or several steps.",
        "answer": "Regular Claude is best for conversation, thinking, drafting, and quick analysis. Cowork is better for longer tasks that involve gathering material, working through a plan, using connected tools, and producing a finished deliverable.",
        "points": [
            "Choose regular chat when you want to stay in the loop turn by turn.",
            "Choose Cowork when the task has a clear result, several steps, and source files or connected tools.",
            "Define the allowed tools, forbidden actions, review points, and final format before Cowork starts.",
            "Keep sensitive or irreversible actions—sending, publishing, deleting, purchasing, or permission changes—behind your approval.",
        ],
        "example": "Ask regular Claude to help shape the outline for a customer report. Use Cowork to collect the approved inputs, build the report, check it against your format, and hand you the draft for review.",
        "screen": ["Chat = think with me", "Cowork = complete this workflow", "Define the finish line"],
        "close": "If you cannot describe the finished deliverable and the stopping point, the task is not ready for Cowork yet.",
        "caption": "Regular Claude or Cowork? Use chat when you want an answer, draft, or thinking partner. Use Cowork when you have a bounded multi-step task, source files, connected tools, and a clear deliverable. Define approvals and the stopping point before it starts.",
        "library": "Regular Claude supports interactive thinking and drafting. Cowork is better for clearly bounded, multi-step work that uses files or tools and ends in a reviewable deliverable.",
    },
    {
        "number": "06",
        "category": "Getting better answers",
        "question": "Why does AI keep giving me generic answers?",
        "hook": "AI gives generic answers when your prompt could have come from anyone.",
        "answer": "Replace broad requests with a real audience, situation, source material, constraints, and a definition of what a good result looks like.",
        "points": [
            "Name the reader and the decision or action the output should support.",
            "Provide examples, notes, customer language, or source documents instead of asking AI to invent context.",
            "Add constraints: length, tone, facts that must remain unchanged, and ideas to avoid.",
            "Ask for a draft, then critique it against explicit criteria before revising.",
        ],
        "example": "Instead of 'write a social post about my landscaping company,' say: 'Write for Raleigh homeowners preparing their yards for summer. Use these three services and this customer question. Keep it under 120 words, practical, and avoid luxury clichés.'",
        "screen": ["Audience", "Situation", "Source material", "Constraints", "Success criteria"],
        "close": "Generic in usually means generic out. Give the AI evidence of what makes your business specific.",
        "caption": "AI sounds generic when your request could belong to anyone. Add the audience, real situation, source material, constraints, and a clear standard for success. The fix is usually better context—not a magical prompt formula.",
        "library": "Generic AI output usually reflects a context-poor request. Improve it by supplying a specific audience, real source material, constraints, examples, and clear evaluation criteria.",
    },
    {
        "number": "07",
        "category": "Getting better answers",
        "question": "How much context should I give Claude or ChatGPT?",
        "hook": "Give AI enough context to make the decision—but not every document your business has ever created.",
        "answer": "Start with the goal, audience, relevant facts, examples, constraints, and desired output. Add more only when the model identifies a real gap or produces the wrong result.",
        "points": [
            "Begin with a short context block: who you are, who the work is for, what you need, and why.",
            "Attach only the source material needed for this task.",
            "Separate facts from instructions so the model knows what must remain true.",
            "Ask the model what information is missing before it drafts.",
        ],
        "example": "For a proposal, include the customer's request, your approved scope, pricing, timeline, and one example proposal. Do not attach your entire Drive and hope the model finds the right truth.",
        "screen": ["Goal", "Audience", "Relevant facts", "Constraints", "Desired output"],
        "close": "A useful test: if a capable new employee would need the information, the AI probably needs it too.",
        "caption": "How much context should you give AI? Enough for a capable new employee to complete the task: goal, audience, relevant facts, examples, constraints, and desired output. Start focused. Add context when a real gap appears—not by dumping your entire business into the prompt.",
        "library": "Effective context is focused rather than exhaustive. Give the model the goal, audience, task-specific facts, examples, constraints, and expected format, then add more only when needed.",
    },
    {
        "number": "08",
        "category": "Getting better answers",
        "question": "How do I get AI to write in my voice?",
        "hook": "Do not describe your voice with five adjectives. Show the AI what you actually sound like.",
        "answer": "Give it several pieces you wrote, explain what each example gets right, define patterns to preserve and avoid, then test the instructions on fresh material.",
        "points": [
            "Use three to five representative samples, not one unusually formal email.",
            "Ask AI to identify sentence length, rhythm, vocabulary, humor, directness, and common transitions.",
            "Correct the analysis: remove habits you do not want repeated and add rules the samples do not reveal.",
            "Save the final voice guide and require a self-check before every draft.",
        ],
        "example": "Tell it: 'I use short openings, concrete examples, and plain language. I do not use hype, exclamation points, or phrases such as game-changer and unlock your potential.' That is more useful than 'sound friendly and professional.'",
        "screen": ["Show 3–5 samples", "Extract patterns", "Correct the guide", "Test on new copy"],
        "close": "Your voice is a pattern the AI can study. Give it evidence, then edit the pattern until it feels true.",
        "caption": "Want AI to write in your voice? Stop giving it adjectives. Share three to five real samples, ask it to identify repeatable patterns, correct the analysis, and save a short voice guide. Include phrases and habits to avoid—not just qualities to imitate.",
        "library": "AI learns a writing voice more reliably from representative samples and explicit corrections than from vague adjectives. Build a short, tested voice guide from actual work.",
    },
    {
        "number": "09",
        "category": "Getting better answers",
        "question": "How do I stop AI from making up information?",
        "hook": "You cannot eliminate AI mistakes with one sentence in a prompt. You reduce them by changing the workflow.",
        "answer": "Ground the task in approved sources, require the model to separate facts from assumptions, demand citations or source references, and keep a human review before the result matters.",
        "points": [
            "Provide the source of truth instead of relying on model memory.",
            "Tell the model to say 'not found' when the source does not support an answer.",
            "Ask it to flag assumptions, uncertainty, conflicting information, and missing fields.",
            "Verify names, dates, prices, links, claims, and decisions before publishing or acting.",
        ],
        "example": "For a property listing, give the approved MLS details and inspection notes. Instruct the model not to infer renovations, school quality, neighborhood safety, or amenities that are absent from the sources.",
        "screen": ["Use approved sources", "Allow 'not found'", "Flag uncertainty", "Verify critical facts"],
        "close": "The goal is not blind trust. The goal is a process that makes unsupported claims easy to spot before they cause damage.",
        "caption": "You cannot prompt hallucinations away. Ground AI in approved sources, allow it to say 'not found,' require uncertainty and assumptions to be labeled, and verify critical facts before publishing or acting. Reliability comes from the workflow.",
        "library": "Reduce invented information by grounding the task in approved sources, making uncertainty explicit, and verifying critical facts. Treat AI output as a draft that must remain traceable to evidence.",
    },
    {
        "number": "10",
        "category": "Getting better answers",
        "question": "Should I create a master prompt for my business?",
        "hook": "A giant master prompt usually becomes a junk drawer: everything is inside, but nobody can find what matters.",
        "answer": "Create a short business context document, then build smaller task-specific instructions for recurring workflows. Keep facts, voice, and process separate so each can be updated without rewriting everything.",
        "points": [
            "Business context: what you do, who you serve, offers, positioning, and facts that rarely change.",
            "Voice guide: how you communicate, examples, preferred language, and phrases to avoid.",
            "Workflow instructions: the steps, inputs, output format, approval rules, and quality checklist for one task.",
            "Review these pieces regularly and version important changes.",
        ],
        "example": "Your listing-description workflow should not carry every rule for invoice follow-up or meeting preparation. Let both workflows reference the same business and voice context while keeping their steps separate.",
        "screen": ["Business context", "Voice guide", "Task instructions", "Quality checklist"],
        "close": "Build a small instruction system, not one enormous prompt. It is easier to test, maintain, and trust.",
        "caption": "Should you create one master prompt for your business? Usually no. Keep a short business context, a separate voice guide, and task-specific workflow instructions. Small, modular instructions are easier to test, update, and reuse than one giant prompt.",
        "library": "Replace a giant master prompt with modular instructions: stable business context, a tested voice guide, and separate workflows for recurring tasks. This structure is easier to maintain and verify.",
    },
]


def set_cell_or_para_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_numbering_definition(document):
    numbering = document.part.numbering_part.element
    abstract = numbering.xpath('./w:abstractNum[@w:abstractNumId="8"]')[0]
    level = abstract.find(qn("w:lvl"))
    p_pr = level.find(qn("w:pPr"))
    tabs = p_pr.find(qn("w:tabs"))
    tab = tabs.find(qn("w:tab"))
    tab.set(qn("w:pos"), "540")
    ind = p_pr.find(qn("w:ind"))
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    return 1


def add_bullet(document, text, num_id):
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return paragraph


def add_label_paragraph(document, label, text, fill=None, after=6):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    if fill:
        set_cell_or_para_shading(paragraph, fill)
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.right_indent = Inches(0.12)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=9, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    return paragraph


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor.from_string(BLUE)
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor.from_string(BLUE)
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(7)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    heading_3.paragraph_format.space_before = Pt(10)
    heading_3.paragraph_format.space_after = Pt(5)
    heading_3.paragraph_format.keep_with_next = True

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    list_bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    list_bullet.font.size = Pt(10.5)
    list_bullet.font.color.rgb = RGBColor.from_string(INK)
    list_bullet.paragraph_format.left_indent = Inches(0.375)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.188)
    list_bullet.paragraph_format.space_before = Pt(0)
    list_bullet.paragraph_format.space_after = Pt(4)
    list_bullet.paragraph_format.line_spacing = 1.25


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    header = section.header
    header.is_linked_to_previous = True
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.paragraph_format.space_after = Pt(0)
    left = header_p.add_run("TWO-MINUTE AI ANSWERS")
    set_run_font(left, size=8, color=MUTED, bold=True)
    right = header_p.add_run("   •   RECORDING GUIDE")
    set_run_font(right, size=8, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = True
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    label = footer_p.add_run("Danny Sullivan  |  ")
    set_run_font(label, size=8, color=MUTED)
    add_page_field(footer_p)


def add_cover(document):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(52)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("CONTENT PRODUCTION GUIDE")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Two-Minute AI Answers")
    set_run_font(run, size=30, color=DARK_BLUE, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("Talking points, captions, and library copy for the first 10 videos")
    set_run_font(run, size=15, color=MUTED)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.left_indent = Inches(0.65)
    note.paragraph_format.right_indent = Inches(0.65)
    note.paragraph_format.space_after = Pt(28)
    set_cell_or_para_shading(note, PALE_LIME)
    run = note.add_run("RECORDING RULE  ")
    set_run_font(run, size=9, color=DARK_BLUE, bold=True)
    run = note.add_run("Lead with the answer. Use one example. End with one next step. Keep each final cut between 90 and 130 seconds.")
    set_run_font(run, size=11, color=INK)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run("Prepared for Danny Sullivan")
    set_run_font(run, size=10.5, color=INK, bold=True)
    meta2 = document.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("Instagram + raleighaiguy.com answer library  |  July 2026")
    set_run_font(run, size=9.5, color=MUTED)

    document.add_page_break()


def add_recording_notes(document, num_id):
    document.add_heading("Use this guide", level=1)
    intro = document.add_paragraph(
        "Each entry is designed as a flexible two-minute talk track rather than a word-for-word script. Keep the phrasing conversational, add a current screen recording or example when useful, and preserve the same core answer across Instagram and the website library."
    )
    intro.paragraph_format.space_after = Pt(10)

    document.add_heading("Recommended video structure", level=2)
    for text in [
        "0:00-0:10 — Hook: name the mistake or tension.",
        "0:10-0:25 — Direct answer: answer the title question before explaining.",
        "0:25-1:20 — Three or four talking points.",
        "1:20-1:45 — One concrete example or miniature demonstration.",
        "1:45-2:00 — One next step and invitation to submit a question.",
    ]:
        add_bullet(document, text, num_id)

    document.add_heading("Production checklist", level=2)
    for text in [
        "Record vertically in 9:16 with your face and captions inside the safe center area.",
        "Put the exact question on screen in the first two seconds.",
        "Use large captions and highlight only the few words that carry the point.",
        "Cut pauses, repeated setup, and tool-history explanations.",
        "Check product names and interface details immediately before recording because AI products change quickly.",
        "End with: “Send me your biggest AI question and I’ll record a two-minute answer.”",
    ]:
        add_bullet(document, text, num_id)

    document.add_page_break()


def add_video(document, video, num_id, is_last=False):
    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(4)
    run = eyebrow.add_run(f"VIDEO {video['number']}  •  {video['category'].upper()}")
    set_run_font(run, size=9, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(video["question"])
    set_run_font(run, size=22, color=DARK_BLUE, bold=True)

    add_label_paragraph(document, "HOOK · 0:00–0:10", video["hook"], fill=PALE_ORANGE, after=8)
    add_label_paragraph(document, "DIRECT ANSWER · 0:10–0:25", video["answer"], fill=PALE_LIME, after=8)

    document.add_heading("Core talking points", level=2)
    for point in video["points"]:
        add_bullet(document, point, num_id)

    add_label_paragraph(document, "EXAMPLE · 1:20–1:45", video["example"], fill=PALE_BLUE, after=8)
    add_label_paragraph(document, "CLOSE · 1:45–2:00", video["close"], after=8)

    document.add_heading("On-screen text", level=3)
    screen = document.add_paragraph()
    screen.paragraph_format.space_after = Pt(7)
    run = screen.add_run("  →  ".join(video["screen"]))
    set_run_font(run, size=10, color=INK, bold=True)

    document.add_heading("Instagram caption", level=3)
    caption = document.add_paragraph(video["caption"])
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.line_spacing = 1.15

    document.add_heading("Library summary", level=3)
    summary = document.add_paragraph(video["library"])
    summary.paragraph_format.space_after = Pt(0)
    summary.paragraph_format.line_spacing = 1.15

    if not is_last:
        document.add_page_break()


def build():
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])
    num_id = add_numbering_definition(document)

    properties = document.core_properties
    properties.title = "Two-Minute AI Answers: First 10 Video Talking Points"
    properties.subject = "Instagram and AI answer library production guide"
    properties.author = "Danny Sullivan"
    properties.keywords = "AI, video, Instagram, talking points, content library"

    add_cover(document)
    add_recording_notes(document, num_id)
    for index, video in enumerate(VIDEOS):
        add_video(document, video, num_id, is_last=index == len(VIDEOS) - 1)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
